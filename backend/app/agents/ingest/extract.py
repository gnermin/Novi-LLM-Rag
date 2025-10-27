from pathlib import Path
from typing import List
import PyPDF2
from docx import Document as DocxDocument
import pandas as pd
from PIL import Image
import pytesseract
import mimetypes
from .base import IngestAgent
from .types import IngestContext, TextBlock, TableData


class ExtractAgent(IngestAgent):
    """
    ExtractAgent - Ekstraktuje sirovi tekst, blokove i tabele iz dokumenata.
    Podržava: PDF, DOCX, XLSX, CSV, slike (OCR)
    """
    
    def __init__(self, ocr_enabled: bool = True):
        super().__init__("ExtractAgent", dependencies=[])
        self.ocr_enabled = ocr_enabled
    
    async def process(self, context: IngestContext):
        """Glavni processing - detektuje tip i ekstraktuje sadržaj"""
        file_path = Path(context.file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"Fajl ne postoji: {context.file_path}")
        
        # Detect file type
        mime_type, _ = mimetypes.guess_type(str(file_path))
        file_ext = file_path.suffix.lower()
        
        # Extract based on type
        if file_ext == '.pdf' or (mime_type and 'pdf' in mime_type):
            await self._extract_pdf(file_path, context)
        
        elif file_ext in ['.docx', '.doc'] or (mime_type and 'word' in mime_type):
            await self._extract_docx(file_path, context)
        
        elif file_ext in ['.xlsx', '.xls']:
            await self._extract_excel(file_path, context)
        
        elif file_ext == '.csv':
            await self._extract_csv(file_path, context)
        
        elif file_ext in ['.png', '.jpg', '.jpeg', '.tiff', '.bmp'] or (mime_type and 'image' in mime_type):
            await self._extract_image(file_path, context)
        
        elif file_ext == '.txt':
            await self._extract_text(file_path, context)
        
        else:
            # Fallback - pokušaj kao plain text
            await self._extract_text(file_path, context)
        
        # Combine blocks into raw_text
        context.raw_text = "\n\n".join([block.text for block in context.blocks])
        
        # Metrics
        context.set_metric("extracted_blocks", len(context.blocks))
        context.set_metric("extracted_tables", len(context.tables))
        context.set_metric("raw_text_length", len(context.raw_text))
    
    async def _extract_pdf(self, file_path: Path, context: IngestContext):
        """Ekstraktuj tekst i blokove iz PDF-a"""
        try:
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    text = page.extract_text()
                    
                    if text.strip():
                        # Split into paragraphs (basic heuristic)
                        paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
                        
                        for para in paragraphs:
                            block = TextBlock(
                                text=para,
                                page=page_num + 1,
                                block_type="paragraph"
                            )
                            context.blocks.append(block)
                            
        except Exception as e:
            raise Exception(f"PDF extraction greška: {str(e)}")
    
    async def _extract_docx(self, file_path: Path, context: IngestContext):
        """Ekstraktuj tekst iz DOCX-a"""
        try:
            doc = DocxDocument(file_path)
            
            for para in doc.paragraphs:
                if para.text.strip():
                    # Detect heading vs paragraph
                    style_name = para.style.name if para.style and hasattr(para.style, 'name') else ""
                    block_type = "heading" if style_name and style_name.startswith('Heading') else "paragraph"
                    
                    block = TextBlock(
                        text=para.text.strip(),
                        block_type=block_type
                    )
                    context.blocks.append(block)
            
            # Extract tables
            for table_idx, table in enumerate(doc.tables):
                headers = [cell.text.strip() for cell in table.rows[0].cells]
                rows = []
                
                for row in table.rows[1:]:
                    row_data = [cell.text.strip() for cell in row.cells]
                    rows.append(row_data)
                
                table_data = TableData(
                    headers=headers,
                    rows=rows,
                    format="csv",
                    metadata={"table_index": table_idx}
                )
                context.tables.append(table_data)
                
        except Exception as e:
            raise Exception(f"DOCX extraction greška: {str(e)}")
    
    async def _extract_excel(self, file_path: Path, context: IngestContext):
        """Ekstraktuj podatke iz Excel-a"""
        try:
            df = pd.read_excel(file_path)
            
            # Store as table
            table_data = TableData(
                headers=df.columns.tolist(),
                rows=df.values.tolist(),
                format="csv",
                metadata={"rows": len(df), "columns": len(df.columns)}
            )
            context.tables.append(table_data)
            
            # Also create text representation
            text = df.to_string(index=False)
            block = TextBlock(
                text=text,
                block_type="table"
            )
            context.blocks.append(block)
            
        except Exception as e:
            raise Exception(f"Excel extraction greška: {str(e)}")
    
    async def _extract_csv(self, file_path: Path, context: IngestContext):
        """Ekstraktuj podatke iz CSV-a"""
        try:
            df = pd.read_csv(file_path)
            
            # Store as table
            table_data = TableData(
                headers=df.columns.tolist(),
                rows=df.values.tolist(),
                format="csv",
                metadata={"rows": len(df), "columns": len(df.columns)}
            )
            context.tables.append(table_data)
            
            # Text representation
            text = df.to_string(index=False)
            block = TextBlock(
                text=text,
                block_type="table"
            )
            context.blocks.append(block)
            
        except Exception as e:
            raise Exception(f"CSV extraction greška: {str(e)}")
    
    async def _extract_image(self, file_path: Path, context: IngestContext):
        """Ekstraktuj tekst iz slike pomoću OCR-a"""
        if not self.ocr_enabled:
            context.add_error("OCR je onemogućen")
            return
        
        try:
            image = Image.open(file_path)
            text = pytesseract.image_to_string(image, lang='bos+eng')
            
            if text.strip():
                block = TextBlock(
                    text=text.strip(),
                    block_type="ocr",
                    confidence=0.8  # Default OCR confidence
                )
                context.blocks.append(block)
                
        except Exception as e:
            context.add_error(f"OCR greška: {str(e)}")
    
    async def _extract_text(self, file_path: Path, context: IngestContext):
        """Ekstraktuj plain text"""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
            
            if text.strip():
                block = TextBlock(
                    text=text.strip(),
                    block_type="text"
                )
                context.blocks.append(block)
                
        except Exception as e:
            raise Exception(f"Text extraction greška: {str(e)}")
